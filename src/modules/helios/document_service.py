"""
Document Generation Service for Helios Study Plans

This service generates DOCX and PDF documents from study plan data using a single
DOCX template as the source of truth for both formats.
"""

import os
import tempfile
import logging
from typing import Dict, Any, List, Optional, BinaryIO
from datetime import datetime, date, timedelta
from pathlib import Path
import base64
from io import BytesIO

from docxtpl import DocxTemplate
from docx2pdf import convert
import json
from jinja2 import Environment, FileSystemLoader, select_autoescape
from weasyprint import HTML, CSS
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.graph_objects as go
import plotly.express as px
from plotly.offline import plot
import numpy as np

from .schemas import StudyPlanSchema, BlockSchema, CycleSchema

logger = logging.getLogger(__name__)


class DocumentGenerationService:
    """Service for generating study plan documents in DOCX and PDF formats."""
    
    def __init__(self, template_path: Optional[str] = None):
        """
        Initialize the document generation service.
        
        Args:
            template_path: Path to the DOCX template file. If None, uses default template.
        """
        self.template_path = template_path or self._get_default_template_path()
        self.pdf_templates_dir = self._get_pdf_templates_dir()
        self.jinja_env = Environment(
            loader=FileSystemLoader(self.pdf_templates_dir),
            autoescape=select_autoescape(['html', 'xml'])
        )
        
        # Set up matplotlib style
        plt.style.use('seaborn-v0_8-whitegrid' if 'seaborn-v0_8-whitegrid' in plt.style.available else 'default')
        sns.set_palette("husl")
        
    def _get_default_template_path(self) -> str:
        """Get the default template path."""
        # Look for template in the assets directory
        base_dir = Path(__file__).parent.parent.parent
        template_path = base_dir / "assets" / "templates" / "study_plan_template.docx"
        return str(template_path)
    
    def _get_pdf_templates_dir(self) -> str:
        """Get the PDF templates directory."""
        base_dir = Path(__file__).parent.parent.parent
        templates_dir = base_dir / "templates" / "pdf_templates"
        return str(templates_dir)
    
    def generate_docx(self, study_plan: StudyPlanSchema, output_path: Optional[str] = None) -> str:
        """
        Generate a DOCX document from study plan data.
        
        Args:
            study_plan: Study plan data to include in the document
            output_path: Optional path to save the document. If None, uses temp file.
            
        Returns:
            Path to the generated DOCX file
            
        Raises:
            FileNotFoundError: If template file is not found
            Exception: If document generation fails
        """
        try:
            # Check if template exists
            if not os.path.exists(self.template_path):
                raise FileNotFoundError(f"Template file not found: {self.template_path}")
            
            # Load the template
            doc = DocxTemplate(self.template_path)
            
            # Prepare context data for template
            context = self._prepare_template_context(study_plan)
            
            # Render the template with context
            doc.render(context)
            
            # Determine output path
            if output_path is None:
                output_path = self._generate_temp_filename(study_plan.study_plan_id, "docx")
            
            # Save the document
            doc.save(output_path)
            
            logger.info(f"Successfully generated DOCX document: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating DOCX document: {e}")
            raise
    
    def generate_pdf(self, study_plan: StudyPlanSchema, output_path: Optional[str] = None) -> str:
        """
        Generate a PDF document from study plan data by first creating DOCX then converting.
        
        Args:
            study_plan: Study plan data to include in the document
            output_path: Optional path to save the PDF. If None, uses temp file.
            
        Returns:
            Path to the generated PDF file
            
        Raises:
            Exception: If document generation or conversion fails
        """
        try:
            # First generate DOCX
            temp_docx_path = self.generate_docx(study_plan)
            
            # Determine PDF output path
            if output_path is None:
                output_path = self._generate_temp_filename(study_plan.study_plan_id, "pdf")
            
            # Convert DOCX to PDF
            convert(temp_docx_path, output_path)
            
            # Clean up temporary DOCX file
            try:
                os.unlink(temp_docx_path)
            except OSError:
                logger.warning(f"Could not delete temporary DOCX file: {temp_docx_path}")
            
            logger.info(f"Successfully generated PDF document: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating PDF document: {e}")
            raise
    
    def generate_enhanced_pdf_overall(self, study_plan: StudyPlanSchema, output_path: Optional[str] = None) -> str:
        """
        Generate a beautiful PDF document with visualizations for the overall study plan.
        
        Args:
            study_plan: Study plan data to include in the document
            output_path: Optional path to save the PDF. If None, uses temp file.
            
        Returns:
            Path to the generated PDF file
        """
        try:
            # Determine output path
            if output_path is None:
                output_path = self._generate_temp_filename(study_plan.study_plan_id, "pdf")
            
            # Prepare context with enhanced data and visualizations
            context = self._prepare_enhanced_template_context(study_plan)
            
            # Generate charts and add them to context
            charts = self._generate_study_plan_charts(study_plan)
            context.update(charts)
            
            # Load and render the HTML template
            template = self.jinja_env.get_template('study_plan_overall.html')
            html_content = template.render(context)
            
            # Generate PDF using WeasyPrint
            HTML(string=html_content, base_url=self.pdf_templates_dir).write_pdf(output_path)
            
            logger.info(f"Successfully generated enhanced PDF overall plan: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating enhanced PDF overall plan: {e}")
            raise
    
    def generate_enhanced_pdf_weekly(self, study_plan: StudyPlanSchema, week_number: int = 1, 
                                   weekly_data: Optional[Dict[str, Any]] = None, 
                                   output_path: Optional[str] = None) -> str:
        """
        Generate a beautiful PDF document for weekly schedule with detailed visualizations.
        
        Args:
            study_plan: Study plan data
            week_number: Week number to generate schedule for
            weekly_data: Optional specific weekly data and schedule
            output_path: Optional path to save the PDF. If None, uses temp file.
            
        Returns:
            Path to the generated PDF file
        """
        try:
            # Determine output path
            if output_path is None:
                output_path = self._generate_temp_filename(f"{study_plan.study_plan_id}_week_{week_number}", "pdf")
            
            # Prepare context for weekly schedule
            context = self._prepare_weekly_template_context(study_plan, week_number, weekly_data)
            
            # Generate weekly charts
            weekly_charts = self._generate_weekly_charts(study_plan, weekly_data)
            context.update(weekly_charts)
            
            # Load and render the HTML template
            template = self.jinja_env.get_template('weekly_schedule.html')
            html_content = template.render(context)
            
            # Generate PDF using WeasyPrint
            HTML(string=html_content, base_url=self.pdf_templates_dir).write_pdf(output_path)
            
            logger.info(f"Successfully generated enhanced PDF weekly schedule: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating enhanced PDF weekly schedule: {e}")
            raise
    
    def _prepare_template_context(self, study_plan: StudyPlanSchema) -> Dict[str, Any]:
        """
        Prepare context data for the document template.
        
        Args:
            study_plan: Study plan data
            
        Returns:
            Dictionary containing template context data
        """
        # Extract all blocks from cycles
        all_blocks = self._extract_blocks_from_cycles(study_plan)
        
        # Group blocks by cycle for better organization
        cycles = self._group_blocks_by_cycle(all_blocks)
        
        # Calculate plan statistics
        stats = self._calculate_plan_statistics(all_blocks)
        
        context = {
            # Basic plan information
            "plan_title": study_plan.title,
            "user_id": study_plan.user_id,
            "generation_date": datetime.now().strftime("%B %d, %Y"),
            "target_year": study_plan.created_for_target_year or "2025",
            "season_context": study_plan.effective_season_context or "General",
            
            # Plan statistics
            "total_blocks": stats["total_blocks"],
            "total_weeks": stats["total_weeks"],
            "total_subjects": stats["total_subjects"],
            "cycles_count": len(cycles),
            
            # Organized content
            "cycles": cycles,
            "blocks": all_blocks,
            
            # Resources
            "curated_resources": study_plan.curated_resources,
            
            # Formatting helpers
            "format_subjects": self._format_subjects_list,
            "format_duration": self._format_duration,
        }
        
        return context
    
    def _prepare_enhanced_template_context(self, study_plan: StudyPlanSchema) -> Dict[str, Any]:
        """
        Prepare enhanced context data for the PDF templates with additional formatting and data.
        
        Args:
            study_plan: Study plan data
            
        Returns:
            Dictionary containing enhanced template context data
        """
        # Get base context
        context = self._prepare_template_context(study_plan)
        
        # Add enhanced formatting and additional computed data
        all_blocks = self._extract_blocks_from_cycles(study_plan)
        stats = self._calculate_plan_statistics(all_blocks)
        
        # Subject distribution analysis
        subject_hours = self._calculate_subject_hours(all_blocks)
        subject_distribution = self._calculate_subject_distribution(subject_hours)
        
        # Weekly breakdown
        weekly_breakdown = self._calculate_weekly_breakdown(all_blocks)
        
        # Enhanced context
        enhanced_context = {
            **context,
            "subject_hours": subject_hours,
            "subject_distribution": subject_distribution,
            "weekly_breakdown": weekly_breakdown,
            "subject_list": stats.get("subject_list", []),
            "avg_hours_per_week": round(sum(subject_hours.values()) / max(context.get("total_weeks", 1), 1), 1),
            "most_intensive_week": max(weekly_breakdown, key=lambda x: x.get("total_hours", 0)) if weekly_breakdown else None,
            "study_intensity": self._calculate_study_intensity(all_blocks),
        }
        
        return enhanced_context
    
    def _prepare_weekly_template_context(self, study_plan: StudyPlanSchema, week_number: int, 
                                       weekly_data: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        Prepare context data for weekly schedule template.
        
        Args:
            study_plan: Study plan data
            week_number: Week number
            weekly_data: Optional weekly schedule data
            
        Returns:
            Dictionary containing weekly template context data
        """
        base_context = self._prepare_template_context(study_plan)
        
        # Generate weekly schedule if not provided
        if not weekly_data:
            weekly_data = self._generate_default_weekly_schedule(study_plan, week_number)
        
        # Calculate week dates
        week_start_date = datetime.now() + timedelta(weeks=week_number - 1)
        day_dates = [(week_start_date + timedelta(days=i)).strftime("%B %d") for i in range(7)]
        
        weekly_context = {
            **base_context,
            "current_week": week_number,
            "week_start_date": week_start_date.strftime("%B %d, %Y"),
            "day_dates": day_dates,
            "weekly_plan": weekly_data.get("schedule", {}),
            "study_days_count": len([d for d in weekly_data.get("schedule", {}).values() if d]),
            "total_weekly_hours": weekly_data.get("total_hours", 40),
            "active_subjects_count": len(weekly_data.get("active_subjects", [])),
            "total_study_hours": weekly_data.get("study_hours", 35),
            "subjects_covered": len(weekly_data.get("active_subjects", [])),
            "practice_sessions": weekly_data.get("practice_sessions", 14),
            "foundation_progress": weekly_data.get("progress", {}).get("foundation", 75),
            "revision_progress": weekly_data.get("progress", {}).get("revision", 60),
            "practice_progress": weekly_data.get("progress", {}).get("practice", 45),
            "weekly_goals": weekly_data.get("goals", []),
            "weekly_notes": weekly_data.get("notes", ""),
        }
        
        return weekly_context
    
    def _group_blocks_by_cycle(self, blocks: List[BlockSchema]) -> List[Dict[str, Any]]:
        """
        Group blocks by their cycle information.
        
        Args:
            blocks: List of blocks to group
            
        Returns:
            List of cycle dictionaries with their associated blocks
        """
        cycles_dict = {}
        uncategorized_blocks = []
        
        for block in blocks:
            if block.cycle_id and block.cycle_order is not None:
                cycle_key = (block.cycle_id, block.cycle_order)
                
                if cycle_key not in cycles_dict:
                    cycles_dict[cycle_key] = {
                        "cycle_id": block.cycle_id,
                        "cycle_name": block.cycle_name or f"Cycle {block.cycle_order + 1}",
                        "cycle_type": block.cycle_type or "StudyCycle",
                        "cycle_order": block.cycle_order,
                        "blocks": []
                    }
                
                cycles_dict[cycle_key]["blocks"].append(block)
            else:
                uncategorized_blocks.append(block)
        
        # Sort cycles by order
        sorted_cycles = sorted(cycles_dict.values(), key=lambda c: c["cycle_order"])
        
        # Add uncategorized blocks as a separate "cycle" if any exist
        if uncategorized_blocks:
            sorted_cycles.append({
                "cycle_id": "uncategorized",
                "cycle_name": "Additional Blocks",
                "cycle_type": "Uncategorized",
                "cycle_order": 999,
                "blocks": uncategorized_blocks
            })
        
        return sorted_cycles
    
    def _extract_blocks_from_cycles(self, study_plan: StudyPlanSchema) -> List[BlockSchema]:
        """Extract all blocks from cycles in a study plan."""
        all_blocks = []
        
        if study_plan.cycles:
            for cycle in study_plan.cycles:
                if cycle.cycle_blocks:
                    all_blocks.extend(cycle.cycle_blocks)
        
        return all_blocks
    
    def _calculate_plan_statistics(self, blocks: List[BlockSchema]) -> Dict[str, Any]:
        """
        Calculate statistics about the study plan.
        
        Args:
            blocks: List of blocks in the plan
            
        Returns:
            Dictionary containing plan statistics
        """
        total_weeks = sum(block.duration_weeks for block in blocks)
        all_subjects = set()
        
        for block in blocks:
            all_subjects.update(block.subjects)
        
        return {
            "total_blocks": len(blocks),
            "total_weeks": total_weeks,
            "total_subjects": len(all_subjects),
            "subject_list": sorted(list(all_subjects))
        }
    
    def _format_subjects_list(self, subjects: List[str]) -> str:
        """Format a list of subjects for display."""
        if not subjects:
            return "No subjects specified"
        elif len(subjects) == 1:
            return subjects[0]
        elif len(subjects) == 2:
            return f"{subjects[0]} and {subjects[1]}"
        else:
            return f"{', '.join(subjects[:-1])}, and {subjects[-1]}"
    
    def _format_duration(self, weeks: int) -> str:
        """Format duration in weeks for display."""
        if weeks == 1:
            return "1 week"
        else:
            return f"{weeks} weeks"
    
    def _generate_temp_filename(self, plan_id: str, extension: str) -> str:
        """
        Generate a temporary filename for the document.
        
        Args:
            plan_id: Study plan ID
            extension: File extension (docx or pdf)
            
        Returns:
            Path to temporary file
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"study_plan_{plan_id}_{timestamp}.{extension}"
        return os.path.join(tempfile.gettempdir(), filename)
    
    def create_default_template(self, output_path: str) -> None:
        """
        Create a default DOCX template file with placeholders.
        
        Args:
            output_path: Path where to save the template
        """
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title
            title = doc.add_heading('{{ plan_title }}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Basic Information
            doc.add_heading('Plan Overview', level=1)
            
            info_table = doc.add_table(rows=5, cols=2)
            info_table.style = 'Table Grid'
            
            info_data = [
                ('User ID:', '{{ user_id }}'),
                ('Generation Date:', '{{ generation_date }}'),
                ('Target Year:', '{{ target_year }}'),
                ('Season Context:', '{{ season_context }}'),
                ('Total Duration:', '{{ total_weeks }} weeks')
            ]
            
            for i, (label, value) in enumerate(info_data):
                info_table.cell(i, 0).text = label
                info_table.cell(i, 1).text = value
            
            # Statistics
            doc.add_heading('Plan Statistics', level=1)
            
            stats_para = doc.add_paragraph()
            stats_para.add_run('Total Blocks: ').bold = True
            stats_para.add_run('{{ total_blocks }}')
            stats_para.add_run('\nTotal Subjects: ').bold = True
            stats_para.add_run('{{ total_subjects }}')
            stats_para.add_run('\nCycles: ').bold = True
            stats_para.add_run('{{ cycles_count }}')
            
            # Cycles Section
            doc.add_heading('Study Plan by Cycles', level=1)
            
            # Template loop for cycles
            doc.add_paragraph('{% for cycle in cycles %}')
            
            cycle_heading = doc.add_heading('{{ cycle.cycle_name }} ({{ cycle.cycle_type }})', level=2)
            
            doc.add_paragraph('{% for block in cycle.blocks %}')
            
            # Block information
            block_para = doc.add_paragraph()
            block_para.add_run('Block: ').bold = True
            block_para.add_run('{{ block.title }}')
            block_para.add_run('\nDuration: ').bold = True
            block_para.add_run('{{ format_duration(block.duration_weeks) }}')
            block_para.add_run('\nSubjects: ').bold = True
            block_para.add_run('{{ format_subjects(block.subjects) }}')
            
            doc.add_paragraph('{% endfor %}')  # End blocks loop
            doc.add_paragraph('{% endfor %}')  # End cycles loop
            
            # Resources Section
            doc.add_heading('Curated Resources', level=1)
            doc.add_paragraph('{{ curated_resources }}')
            
            # Save the template
            doc.save(output_path)
            logger.info(f"Created default template at: {output_path}")
            
        except Exception as e:
            logger.error(f"Error creating default template: {e}")
            raise
    
    def _generate_study_plan_charts(self, study_plan: StudyPlanSchema) -> Dict[str, str]:
        """
        Generate charts for the study plan overview and return them as base64 encoded strings.
        
        Args:
            study_plan: Study plan data
            
        Returns:
            Dictionary containing chart data as base64 strings
        """
        charts = {}
        
        try:
            all_blocks = self._extract_blocks_from_cycles(study_plan)
            subject_hours = self._calculate_subject_hours(all_blocks)
            
            # Generate subjects pie chart
            if subject_hours:
                charts['subjects_pie_chart'] = self._create_subjects_pie_chart(subject_hours)
            
            # Generate timeline chart
            charts['timeline_chart'] = self._create_timeline_chart(all_blocks)
            
            # Generate weekly distribution chart
            weekly_data = self._calculate_weekly_breakdown(all_blocks)
            if weekly_data:
                charts['weekly_distribution_chart'] = self._create_weekly_distribution_chart(weekly_data)
                
        except Exception as e:
            logger.warning(f"Error generating charts: {e}")
            
        return charts
    
    def _generate_weekly_charts(self, study_plan: StudyPlanSchema, weekly_data: Optional[Dict[str, Any]] = None) -> Dict[str, str]:
        """
        Generate charts for weekly schedule.
        
        Args:
            study_plan: Study plan data
            weekly_data: Weekly schedule data
            
        Returns:
            Dictionary containing chart data as base64 strings
        """
        charts = {}
        
        try:
            if not weekly_data:
                weekly_data = self._generate_default_weekly_schedule(study_plan, 1)
                
            # Generate daily time distribution chart
            charts['daily_time_chart'] = self._create_daily_time_distribution_chart(weekly_data)
            
            # Generate subject balance chart
            charts['subject_balance_chart'] = self._create_subject_balance_chart(weekly_data)
                
        except Exception as e:
            logger.warning(f"Error generating weekly charts: {e}")
            
        return charts
    
    def _create_subjects_pie_chart(self, subject_hours: Dict[str, float]) -> str:
        """Create a pie chart for subject time distribution."""
        try:
            plt.figure(figsize=(10, 8))
            
            # Create pie chart with custom colors
            colors = plt.cm.Set3(np.linspace(0, 1, len(subject_hours)))
            wedges, texts, autotexts = plt.pie(
                subject_hours.values(), 
                labels=subject_hours.keys(),
                autopct='%1.1f%%',
                colors=colors,
                startangle=90,
                textprops={'fontsize': 12, 'weight': 'bold'}
            )
            
            plt.title('Study Time Distribution by Subject', fontsize=16, weight='bold', pad=20)
            
            # Save to base64
            return self._save_plot_to_base64()
            
        except Exception as e:
            logger.error(f"Error creating subjects pie chart: {e}")
            return ""
    
    def _create_timeline_chart(self, blocks: List[BlockSchema]) -> str:
        """Create a timeline chart showing study blocks progression."""
        try:
            if not blocks:
                return ""
                
            plt.figure(figsize=(14, 8))
            
            # Prepare data
            block_names = [block.title for block in blocks]
            durations = [block.duration_weeks for block in blocks]
            subjects = [', '.join(block.subjects[:2]) + ('...' if len(block.subjects) > 2 else '') for block in blocks]
            
            # Create horizontal bar chart
            y_pos = np.arange(len(block_names))
            bars = plt.barh(y_pos, durations, color=plt.cm.viridis(np.linspace(0, 1, len(blocks))))
            
            plt.yticks(y_pos, [f"{name}\n({subject})" for name, subject in zip(block_names, subjects)])
            plt.xlabel('Duration (Weeks)', fontsize=12, weight='bold')
            plt.title('Study Blocks Timeline', fontsize=16, weight='bold')
            
            # Add value labels on bars
            for i, (bar, duration) in enumerate(zip(bars, durations)):
                plt.text(bar.get_width() + 0.1, bar.get_y() + bar.get_height()/2, 
                        f'{duration}w', ha='left', va='center', fontweight='bold')
            
            plt.tight_layout()
            return self._save_plot_to_base64()
            
        except Exception as e:
            logger.error(f"Error creating timeline chart: {e}")
            return ""
    
    def _create_weekly_distribution_chart(self, weekly_data: List[Dict[str, Any]]) -> str:
        """Create a chart showing weekly study load distribution."""
        try:
            plt.figure(figsize=(12, 6))
            
            weeks = [f"Week {i+1}" for i in range(len(weekly_data))]
            hours = [week.get('total_hours', 0) for week in weekly_data]
            
            bars = plt.bar(weeks, hours, color=plt.cm.plasma(np.linspace(0, 1, len(weeks))))
            
            plt.xlabel('Week', fontsize=12, weight='bold')
            plt.ylabel('Total Hours', fontsize=12, weight='bold')
            plt.title('Weekly Study Load Distribution', fontsize=16, weight='bold')
            plt.xticks(rotation=45)
            
            # Add value labels on bars
            for bar, hour in zip(bars, hours):
                plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
                        f'{hour}h', ha='center', va='bottom', fontweight='bold')
            
            plt.tight_layout()
            return self._save_plot_to_base64()
            
        except Exception as e:
            logger.error(f"Error creating weekly distribution chart: {e}")
            return ""
    
    def _create_daily_time_distribution_chart(self, weekly_data: Dict[str, Any]) -> str:
        """Create a chart showing daily time distribution for the week."""
        try:
            plt.figure(figsize=(12, 6))
            
            # Sample daily hours (this would come from actual schedule data)
            days = ['Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat', 'Sun']
            study_hours = [6, 7, 6, 8, 7, 5, 4]  # This should come from actual data
            
            bars = plt.bar(days, study_hours, color=plt.cm.coolwarm(np.linspace(0, 1, 7)))
            
            plt.xlabel('Day of Week', fontsize=12, weight='bold')
            plt.ylabel('Study Hours', fontsize=12, weight='bold')
            plt.title('Daily Study Hours Distribution', fontsize=16, weight='bold')
            
            # Add value labels on bars
            for bar, hour in zip(bars, study_hours):
                plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1,
                        f'{hour}h', ha='center', va='bottom', fontweight='bold')
            
            plt.tight_layout()
            return self._save_plot_to_base64()
            
        except Exception as e:
            logger.error(f"Error creating daily time distribution chart: {e}")
            return ""
    
    def _create_subject_balance_chart(self, weekly_data: Dict[str, Any]) -> str:
        """Create a chart showing subject balance for the week."""
        try:
            # Sample subject data (this should come from actual schedule)
            subjects = ['Math', 'Physics', 'Chemistry', 'Biology']
            hours = [8, 7, 6, 5]
            
            plt.figure(figsize=(10, 8))
            colors = plt.cm.Set2(np.linspace(0, 1, len(subjects)))
            
            wedges, texts, autotexts = plt.pie(
                hours, 
                labels=subjects,
                autopct='%1.1f%%',
                colors=colors,
                startangle=90,
                textprops={'fontsize': 12, 'weight': 'bold'}
            )
            
            plt.title('Weekly Subject Time Balance', fontsize=16, weight='bold', pad=20)
            return self._save_plot_to_base64()
            
        except Exception as e:
            logger.error(f"Error creating subject balance chart: {e}")
            return ""
    
    def _save_plot_to_base64(self) -> str:
        """Save current matplotlib plot to base64 string."""
        try:
            buffer = BytesIO()
            plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight', 
                       facecolor='white', edgecolor='none')
            buffer.seek(0)
            image_base64 = base64.b64encode(buffer.getvalue()).decode()
            plt.close()  # Important: close the figure to free memory
            return f"data:image/png;base64,{image_base64}"
        except Exception as e:
            logger.error(f"Error saving plot to base64: {e}")
            plt.close()
            return ""
    
    def _calculate_subject_hours(self, blocks: List[BlockSchema]) -> Dict[str, float]:
        """Calculate total hours per subject across all blocks."""
        subject_hours = {}
        
        for block in blocks:
            # Estimate hours per week (this could be more sophisticated)
            hours_per_week = 40  # Default assumption
            total_hours = hours_per_week * block.duration_weeks
            hours_per_subject = total_hours / max(len(block.subjects), 1)
            
            for subject in block.subjects:
                subject_hours[subject] = subject_hours.get(subject, 0) + hours_per_subject
                
        return subject_hours
    
    def _calculate_subject_distribution(self, subject_hours: Dict[str, float]) -> Dict[str, float]:
        """Calculate percentage distribution of subjects."""
        total_hours = sum(subject_hours.values())
        if total_hours == 0:
            return {}
            
        return {subject: (hours / total_hours) * 100 
               for subject, hours in subject_hours.items()}
    
    def _calculate_weekly_breakdown(self, blocks: List[BlockSchema]) -> List[Dict[str, Any]]:
        """Calculate weekly breakdown of study load."""
        weekly_breakdown = []
        current_week = 1
        
        for block in blocks:
            for week in range(block.duration_weeks):
                weekly_breakdown.append({
                    "week_number": current_week,
                    "block_title": block.title,
                    "subjects": block.subjects,
                    "total_hours": 40,  # This could be calculated from weekly_plan
                    "focus_areas": block.subjects[:2]  # Top 2 subjects
                })
                current_week += 1
                
        return weekly_breakdown
    
    def _calculate_study_intensity(self, blocks: List[BlockSchema]) -> str:
        """Calculate overall study intensity level."""
        if not blocks:
            return "Low"
            
        total_weeks = sum(block.duration_weeks for block in blocks)
        total_subjects = len(set(subject for block in blocks for subject in block.subjects))
        
        # Simple intensity calculation
        intensity_score = (total_subjects * 2) + (total_weeks * 0.5)
        
        if intensity_score < 10:
            return "Low"
        elif intensity_score < 20:
            return "Medium"
        else:
            return "High"
    
    def _generate_default_weekly_schedule(self, study_plan: StudyPlanSchema, week_number: int) -> Dict[str, Any]:
        """Generate a default weekly schedule structure."""
        all_blocks = self._extract_blocks_from_cycles(study_plan)
        active_subjects = list(set(subject for block in all_blocks for subject in block.subjects))[:4]
        
        return {
            "schedule": {},  # Would contain detailed daily schedules
            "total_hours": 40,
            "study_hours": 35,
            "active_subjects": active_subjects,
            "practice_sessions": 14,
            "progress": {
                "foundation": 75,
                "revision": 60,
                "practice": 45
            },
            "goals": [
                "Complete foundation chapters for all active subjects",
                "Solve practice problems for each topic covered",
                "Review and consolidate previous week's learning",
                "Take at least 2 mock assessments",
                "Identify and address weak areas"
            ],
            "notes": "Remember to take regular breaks, stay hydrated, and maintain a consistent study routine."
        }


# Utility functions for external use
def generate_study_plan_docx(study_plan: StudyPlanSchema, output_path: Optional[str] = None) -> str:
    """
    Convenience function to generate a DOCX document from study plan data.
    
    Args:
        study_plan: Study plan data
        output_path: Optional output path
        
    Returns:
        Path to generated DOCX file
    """
    service = DocumentGenerationService()
    return service.generate_docx(study_plan, output_path)


def generate_study_plan_pdf(study_plan: StudyPlanSchema, output_path: Optional[str] = None) -> str:
    """
    Convenience function to generate a PDF document from study plan data.
    
    Args:
        study_plan: Study plan data
        output_path: Optional output path
        
    Returns:
        Path to generated PDF file
    """
    service = DocumentGenerationService()
    return service.generate_pdf(study_plan, output_path)


def generate_enhanced_study_plan_pdf(study_plan: StudyPlanSchema, output_path: Optional[str] = None) -> str:
    """
    Convenience function to generate an enhanced PDF document with visualizations.
    
    Args:
        study_plan: Study plan data
        output_path: Optional output path
        
    Returns:
        Path to generated enhanced PDF file
    """
    service = DocumentGenerationService()
    return service.generate_enhanced_pdf_overall(study_plan, output_path)


def generate_weekly_schedule_pdf(study_plan: StudyPlanSchema, week_number: int = 1, 
                                weekly_data: Optional[Dict[str, Any]] = None,
                                output_path: Optional[str] = None) -> str:
    """
    Convenience function to generate a weekly schedule PDF with visualizations.
    
    Args:
        study_plan: Study plan data
        week_number: Week number to generate
        weekly_data: Optional weekly schedule data
        output_path: Optional output path
        
    Returns:
        Path to generated weekly schedule PDF file
    """
    service = DocumentGenerationService()
    return service.generate_enhanced_pdf_weekly(study_plan, week_number, weekly_data, output_path)
