#!/usr/bin/env python3
"""
Create a proper DOCX template for study plan generation
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_study_plan_template():
    # Create a new document
    doc = Document()
    
    # Add title
    title = doc.add_heading('{{title}}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add student info section
    doc.add_heading('Student Information', level=1)
    doc.add_paragraph('Student ID: {{user_id}}')
    doc.add_paragraph('Target Year: {{created_for_target_year}}')
    doc.add_paragraph('Plan ID: {{study_plan_id}}')
    
    # Add plan overview
    doc.add_heading('Study Plan Overview', level=1)
    doc.add_paragraph('{{effective_season_context}}')
    
    # Add blocks section
    doc.add_heading('Study Blocks', level=1)
    doc.add_paragraph('{% for block in blocks %}')
    
    # Block template
    doc.add_heading('{{block.title}}', level=2)
    doc.add_paragraph('Block ID: {{block.block_id}}')
    doc.add_paragraph('Duration: {{block.duration_weeks}} weeks')
    
    # Cycle information
    doc.add_paragraph('Cycle: {{block.cycle_name}} ({{block.cycle_type}})')
    doc.add_paragraph('Cycle Order: {{block.cycle_order}}')
    
    # Subjects
    doc.add_paragraph('Subjects:')
    doc.add_paragraph('{% for subject in block.subjects %}')
    doc.add_paragraph('• {{subject}}')
    doc.add_paragraph('{% endfor %}')
    
    # Resources
    doc.add_paragraph('Resources:')
    doc.add_paragraph('{% for category, items in block.resources.items() %}')
    doc.add_paragraph('{{category|title}}:')
    doc.add_paragraph('{% for item in items %}')
    doc.add_paragraph('  • {{item}}')
    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph('{% endfor %}')
    
    # Weekly plan
    doc.add_paragraph('Weekly Plan: {{block.weekly_plan}}')
    
    doc.add_paragraph('{% endfor %}')
    
    # Add resources section
    doc.add_heading('Curated Resources', level=1)
    doc.add_paragraph('{{curated_resources}}')
    
    # Save the template
    template_path = '/laex/upscapp-be/src/assets/templates/study_plan_template.docx'
    doc.save(template_path)
    print(f"Study plan template created at: {template_path}")

if __name__ == "__main__":
    create_study_plan_template()
